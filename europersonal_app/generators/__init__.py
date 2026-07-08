"""Generare documente .docx pe baza șabloanelor (exemplele corecte cu placeholder-e).

Strategia: NU construim documente de la zero. Folosim exemplele corect
formatate ca șabloane, cu placeholder-e {{...}}, și înlocuim textul la
nivel de "run" ca să păstrăm formatarea identică.
"""
from pathlib import Path

import docx
from docx.text.paragraph import Paragraph

# rădăcina aplicației = folderul de deasupra lui generators/
TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"


# ---------------------------------------------------------------- utilitare

def _iter_all_paragraphs(doc):
    def from_table(table):
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for t in cell.tables:
                    yield from from_table(t)
    yield from doc.paragraphs
    for table in doc.tables:
        yield from from_table(table)
    for section in doc.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            for table in part.tables:
                yield from from_table(table)


def _replace_in_paragraph(p: Paragraph, mapping: dict) -> None:
    for key, value in mapping.items():
        for run in p.runs:
            if key in run.text:
                run.text = run.text.replace(key, value)
        if key in p.text:
            runs = p.runs
            full = "".join(r.text for r in runs)
            start = full.find(key)
            pos = 0
            first_idx = None
            for i, r in enumerate(runs):
                nxt = pos + len(r.text)
                if first_idx is None and start < nxt:
                    first_idx = i
                pos = nxt
            new_full = full.replace(key, value)
            prefix = "".join(r.text for r in runs[:first_idx])
            runs[first_idx].text = new_full[len(prefix):]
            for r in runs[first_idx + 1:]:
                r.text = ""


def render_template(template_name: str, mapping: dict) -> docx.Document:
    doc = docx.Document(TEMPLATES_DIR / template_name)
    for p in _iter_all_paragraphs(doc):
        if any(k in p.text for k in mapping):
            _replace_in_paragraph(p, mapping)
    return doc


def check_placeholders_left(doc) -> list:
    return [p.text.strip()[:80] for p in _iter_all_paragraphs(doc) if "{{" in p.text]


def _save(doc, output_path):
    leftovers = check_placeholders_left(doc)
    if leftovers:
        raise ValueError(f"Placeholder-e necompletate: {leftovers}")
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------- documente

def generate_confirmare(package, output_path):
    """CONFIRMARE — autorizare de folosire a spațiului locativ."""
    from models import fmt_date
    mapping = {
        "{{NUME_COMPLET}}": package.passport.full_name,
        "{{ADRESA_CAZARE}}": package.accommodation.address,
        "{{TIP_CONTRACT}}": package.accommodation.contract_type,
        "{{NR_CONTRACT}}": package.accommodation.contract_number,
        "{{DATA_CONTRACT}}": fmt_date(package.accommodation.contract_date),
    }
    return _save(render_template("CONFIRMARE.docx", mapping), output_path)


def confirmare_filename(package):
    return f"CONFIRMARE_{package.passport.full_name.replace(' ', '_')}.docx"


def generate_demers_invitatie(package, output_path):
    """DEMERS INVITATIE — scrisoarea scurtă către Direcția Regională."""
    from models import fmt_date
    if not package.employment or not package.employment.job_title:
        raise ValueError("Lipsește funcția/ocupația lucrătorului (necesară în tabel).")
    if not package.meta.demers_number or not package.meta.demers_date:
        raise ValueError("Lipsesc numărul și data demersului.")
    mapping = {
        "{{NR_DEMERS}}": package.meta.demers_number,
        "{{DATA_DEMERS}}": fmt_date(package.meta.demers_date),
        "{{TARA}}": package.passport.country_ro,
        "{{NUME_COMPLET}}": package.passport.full_name,
        "{{NR_PASAPORT}}": package.passport.passport_number,
        "{{FUNCTIA}}": package.employment.job_title,
    }
    return _save(render_template("DEMERS_INVITATIE.docx", mapping), output_path)


def demers_filename(package):
    return f"DEMERS_INVITATIE_{package.passport.full_name.replace(' ', '_')}.docx"


def generate_anexa1(package, output_path):
    """ANEXA 1 — angajamentul de luare la întreținere și/sau cazare."""
    from models import fmt_date
    if not package.passport.date_of_birth:
        raise ValueError("Lipsește data nașterii (necesară în ANEXA 1).")
    if not package.passport.date_of_expiry:
        raise ValueError("Lipsește valabilitatea pașaportului (necesară în ANEXA 1).")
    mapping = {
        "{{NUME_COMPLET}}": package.passport.full_name,
        "{{DATA_NASTERII}}": fmt_date(package.passport.date_of_birth),
        "{{CETATENIA}}": package.passport.citizenship,
        "{{NR_PASAPORT}}": package.passport.passport_number,
        "{{VALABILITATEA}}": fmt_date(package.passport.date_of_expiry),
        "{{ADRESA_CAZARE}}": package.accommodation.address,
    }
    return _save(render_template("ANEXA_1.docx", mapping), output_path)


def anexa1_filename(package):
    return f"ANEXA_1_{package.passport.full_name.replace(' ', '_')}.docx"


def generate_cim(package, output_path):
    """CIM — Contract Individual de Muncă (bilingv RO/EN)."""
    from models import fmt_date
    if not package.employment or not package.employment.job_title:
        raise ValueError("Lipsește funcția (necesară în CIM).")
    if not package.meta.cim_number or not package.meta.cim_date:
        raise ValueError("Lipsesc numărul și data CIM.")
    emp = package.employment
    job_en = emp.job_title_en.strip() if emp.job_title_en else ""
    if not job_en:
        try:
            from config import JOB_TITLES_EN
            job_en = JOB_TITLES_EN.get(emp.job_title, emp.job_title)
        except ImportError:
            job_en = emp.job_title
    mapping = {
        "{{NR_CIM}}": package.meta.cim_number,
        "{{DATA_CIM}}": fmt_date(package.meta.cim_date),
        "{{NUME_COMPLET}}": package.passport.full_name,
        "{{CETATENIA}}": package.passport.citizenship,
        "{{DATA_NASTERII}}": fmt_date(package.passport.date_of_birth),
        "{{LOCUL_NASTERII}}": package.passport.place_of_birth,
        "{{NR_PASAPORT}}": package.passport.passport_number,
        "{{VALABILITATEA}}": fmt_date(package.passport.date_of_expiry),
        "{{FUNCTIA_RO}}": emp.job_title,
        "{{FUNCTIA_EN}}": job_en,
        "{{SALARIU_RO}}": emp.salary_ro,
        "{{SALARIU_EN}}": emp.salary_en,
    }
    return _save(render_template("CIM.docx", mapping), output_path)


def cim_filename(package):
    return f"CIM_{package.passport.full_name.replace(' ', '_')}.docx"


def generate_invitatie_pj(package, output_path):
    """INVITATIE — DEMERS pentru eliberarea invitației (persoană juridică).

    Nepal: prenume tată „-”, domiciliat linie liberă; seria/nr. separate.
    """
    from models import fmt_date
    p = package.passport
    mapping = {
        "{{NUMELE}}": p.surname,
        "{{PRENUMELE}}": p.given_name,
        "{{PRENUME_TATA}}": p.father_name if p.father_name else "-",
        "{{CETATENIA}}": p.citizenship,
        "{{DATA_NASTERII}}": fmt_date(p.date_of_birth),
        "{{LOCUL_NASTERII}}": p.place_of_birth,
        "{{DOMICILIAT}}": (p.permanent_address if p.permanent_address
                           else "_____________________________________________"),
        "{{SEXUL}}": p.sex,
        "{{SERIA}}": p.passport_series,
        "{{NR_DOC}}": p.passport_no_digits,
        "{{DATA_ELIBERARII}}": fmt_date(p.date_of_issue),
        "{{VALABIL}}": fmt_date(p.date_of_expiry),
        "{{TARA_INTRARE}}": p.country_ro.upper(),
    }
    return _save(render_template("INVITATIE_PJ.docx", mapping), output_path)


def invitatie_pj_filename(package):
    return f"INVITATIE_{package.passport.full_name.replace(' ', '_')}.docx"


# ---------------------------------------------------------------- pachet

AVAILABLE_DOCUMENTS = {
    "CIM (Contract Individual de Muncă)": (generate_cim, cim_filename),
    "DEMERS INVITATIE (scrisoare Direcția Regională)": (generate_demers_invitatie, demers_filename),
    "INVITATIE (persoană juridică)": (generate_invitatie_pj, invitatie_pj_filename),
    "CONFIRMARE": (generate_confirmare, confirmare_filename),
    "ANEXA 1 (opțional)": (generate_anexa1, anexa1_filename),
}

# Pachetul-țintă pentru „Crează dosar nou” (5 documente):
# CIM · DEMERS Model M (urmează) · DEMERS INVITATIE · INVITATIE · CONFIRMARE
PACKAGE_DOCUMENTS = [
    "CIM (Contract Individual de Muncă)",
    # "DEMERS Model M" — se adaugă la implementare
    "DEMERS INVITATIE (scrisoare Direcția Regională)",
    "INVITATIE (persoană juridică)",
    "CONFIRMARE",
]


def generate_package_zip(package, selected):
    import io
    import tempfile
    import zipfile
    filenames = []
    buf = io.BytesIO()
    with tempfile.TemporaryDirectory() as tmp:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for doc_name in selected:
                gen_fn, name_fn = AVAILABLE_DOCUMENTS[doc_name]
                fname = name_fn(package)
                path = gen_fn(package, Path(tmp) / fname)
                zf.write(path, fname)
                filenames.append(fname)
    return buf.getvalue(), filenames
