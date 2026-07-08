"""Extragerea adresei de cazare din Contractul de Comodat / Locațiune.

Adresa din contract este SURSA UNICĂ de adevăr pentru tot pachetul.
Suportă .docx și .pdf (cu strat de text). PDF-urile scanate fără text
returnează listă goală -> utilizatorul introduce adresa manual în UI.

Testat pe variații de formulare: "situată în", "amplasat în", "la adresa:",
ordinea stradă-întâi ("str. X nr. 5, ap. 2, mun. Chișinău"), adrese fără
"nr.", bloc, "ap./ înc.", date cu luni în litere ("5 martie 2025").
"""
import io
import re
import unicodedata

import docx


# ---------------------------------------------------------------- extragere text

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Text complet din .docx sau .pdf."""
    name = filename.lower()
    if name.endswith(".docx"):
        d = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    raise ValueError("Format nesuportat. Încărcați .docx sau .pdf.")


# ---------------------------------------------------------------- căutare adresă
# Componente reutilizabile. Nota bene:
#  - \b previne potriviri false ("amplasat" conține "sat")
#  - (?-i:...) forțează majusculă la inițiala numelui, deși restul e IGNORECASE

_CITY = r"\b(?:mun\.?|or\.?|orașul|municipiul|s\.|sat(?:ul)?|com(?:una)?\.?)\s*(?-i:[A-ZȘȚĂÎÂ])[^,\n;]{2,30}"
_SECTOR = r"(?:sect(?:orul)?\.?\s*(?-i:[A-ZȘȚĂÎÂ])[^,\n;]{2,20})"
_STREET = r"\b(?:str\.?|strada|bd\.?|bulevardul|șos\.?|sos\.?|șoseaua)\s*[^,\n;]{2,45}"
_NRPART = r"(?:\s*,?\s*nr\.?\s*[\w/]{1,8})?"
_BLOC = r"(?:bl(?:oc)?\.?\s*\w{1,5})"
_AP = r"(?:ap(?:artamentul)?\.?\s*(?:/\s*înc\.?\s*)?\w{1,6}|înc(?:ăperea)?\.?\s*\w{1,6})"

# Ordinea 1: oraș, [sector], stradă, [nr], [bloc], [ap]
_ADDRESS_RE_CITY_FIRST = re.compile(
    rf"{_CITY}(?:\s*,\s*{_SECTOR})?\s*,\s*{_STREET}"
    rf"{_NRPART}(?:\s*,\s*{_BLOC})?(?:\s*,\s*{_AP})?",
    re.IGNORECASE,
)

# Ordinea 2: strada întâi, orașul după (ex. "str. București nr. 83, ap. 9, mun. Chișinău")
_ADDRESS_RE_STREET_FIRST = re.compile(
    rf"{_STREET}{_NRPART}(?:\s*,\s*{_BLOC})?"
    rf"(?:\s*,\s*{_AP})?\s*,\s*{_CITY}(?:\s*,\s*{_SECTOR})?",
    re.IGNORECASE,
)

# Cuvinte-cheie care indică adresa DE CAZARE (nu sediul firmei)
_HOUSING_HINTS = (
    "va locui", "vor locui", "cazare", "cazarea", "imobil", "încăperea locativ",
    "incaperea locativ", "apartament", "locuinț", "locuint", "folosință",
    "folosinta", "folosire", "comodat", "locațiune", "locatiune", "chirie",
    "transmite", "obiectul", "situat", "situată", "amplasat", "amplasată",
    "domicilia", "domiciliu", "spațiul locativ", "spatiul locativ",
    "adresa:", "la adresa", "pe adresa",
)

# Sediul firmei — trebuie EXCLUS din candidați
_COMPANY_ADDRESS_MARKERS = ("bănulescu-bodoni", "banulescu-bodoni", "banulescu bodoni")


def _norm(s: str) -> str:
    s = unicodedata.normalize("NFC", s)
    return re.sub(r"\s+", " ", s).strip().rstrip(".").rstrip(",")


def find_address_candidates(text: str) -> list[str]:
    """Adrese candidate, ordonate: cele din context de cazare primele.

    Scor = numărul de indicii de cazare în jurul adresei (+ bonus dacă adresa
    e completă, cu nr. și ap.). Sediul EUROPERSONAL e exclus întotdeauna.
    """
    scored: list[tuple[int, int, str]] = []
    seen = set()
    for order, rx in enumerate((_ADDRESS_RE_CITY_FIRST, _ADDRESS_RE_STREET_FIRST)):
        for m in rx.finditer(text):
            addr = _norm(m.group(0))
            low = addr.lower()
            if any(marker in low for marker in _COMPANY_ADDRESS_MARKERS):
                continue  # sediul EUROPERSONAL, nu cazarea
            if low in seen:
                continue
            seen.add(low)
            context = text[max(0, m.start() - 200): m.end() + 80].lower()
            score = sum(hint in context for hint in _HOUSING_HINTS)
            if re.search(r"nr\.?\s*\d", low):
                score += 1
            if re.search(r"ap\.?|înc\.?", low):
                score += 1
            scored.append((score, -order, addr))
    scored.sort(key=lambda t: (-t[0], t[1]))
    return [a for _, _, a in scored]


# ---------------------------------------------------------------- detalii contract

_MONTHS_RO = {
    "ianuarie": 1, "februarie": 2, "martie": 3, "aprilie": 4, "mai": 5,
    "iunie": 6, "iulie": 7, "august": 8, "septembrie": 9, "octombrie": 10,
    "noiembrie": 11, "decembrie": 12,
}


def _find_date(text: str) -> str | None:
    """Prima dată din text: 31.10.2025, 31/10/2025 sau „5 martie 2025”."""
    m = re.search(r"\b(\d{1,2})[./](\d{1,2})[./](\d{4})\b", text)
    if m:
        return f"{int(m.group(1)):02d}.{int(m.group(2)):02d}.{m.group(3)}"
    m = re.search(rf"\b(\d{{1,2}})\s+({'|'.join(_MONTHS_RO)})\s+(\d{{4}})\b", text, re.I)
    if m:
        return f"{int(m.group(1)):02d}.{_MONTHS_RO[m.group(2).lower()]:02d}.{m.group(3)}"
    return None


def find_contract_details(text: str) -> dict:
    """Tip / nr. / dată contract, cu toleranță la variații de formulare."""
    out = {}
    low = text.lower()
    if "comodat" in low:
        out["contract_type"] = "comodat"
    elif "locațiune" in low or "locatiune" in low or "chirie" in low:
        out["contract_type"] = "locațiune"

    # nr. contract: "nr. F/N din", "Nr. 12/2025 încheiat la", "nr. 5, ..."
    m = re.search(
        r"nr\.?\s*([A-Za-z0-9/\-]{1,12})\s*(?:din\b|încheiat|incheiat|,|la\b)",
        text, re.I)
    if m:
        out["contract_number"] = m.group(1)
    elif re.search(r"\bf\s*/\s*n\b|fără\s+număr|fara\s+numar", low):
        out["contract_number"] = "F/N"

    # data: întâi lângă cuvinte-cheie ("din", "încheiat la data de"), apoi oriunde
    for m in re.finditer(
            r"(?:\bdin\b|încheiat|incheiat|la\s+data\s+de)\s*(?:la\s+)?(?:data\s+de\s+)?(.{0,30})",
            text, re.I):
        d = _find_date(m.group(1))
        if d:
            out["contract_date"] = d
            break
    if "contract_date" not in out:
        d = _find_date(text)
        if d:
            out["contract_date"] = d
    return out


def parse_contract(file_bytes: bytes, filename: str) -> dict:
    """Rezultat complet: adrese candidate + detalii contract."""
    text = extract_text(file_bytes, filename)
    return {
        "text_found": bool(text.strip()),
        "addresses": find_address_candidates(text),
        **find_contract_details(text),
    }
