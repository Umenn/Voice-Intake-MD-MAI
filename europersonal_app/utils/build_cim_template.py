"""Construiește templates/CIM.docx — șablonul bilingv (RO/EN) al Contractului
Individual de Muncă, reconstruit fidel după exemplul K C NABIN (PDF).

Rulare:  python build_cim_template.py
Markup:  **text** = bold în interiorul paragrafelor.
"""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

FONT = "Times New Roman"
SIZE = Pt(10)

def add_markup_text(par, text, base_bold=False):
    for chunk in re.split(r"(\*\*.*?\*\*)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = par.add_run(chunk[2:-2]); r.bold = True
        else:
            r = par.add_run(chunk); r.bold = base_bold
        r.font.name = FONT; r.font.size = SIZE

def fill_cell(cell, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.alignment = align
        p.paragraph_format.space_after = Pt(4)
        add_markup_text(p, line, base_bold=bold)

# (RO, EN) — textul integral, cu placeholder-e {{...}}
S = []
S.append((
"**CONTRACT INDIVIDUAL DE MUNCĂ**\n**Nr. {{NR_CIM}}**\n**Data: {{DATA_CIM}}**\n**Locul: mun. Chișinău**",
"**INDIVIDUAL LABOR CONTRACT**\n**No. {{NR_CIM}}**\n**Date: {{DATA_CIM}}**\n**Place: Chisinau Municipality**"))
S.append((
"""**1. Părțile contractante**
„Europersonal” SRL, cu sediul în mun. Chișinău, str. Mitropolit Gavriil Bănulescu-Bodoni, nr. 43a, ap. 14, c/f 1016600043760, reprezentată prin director administrativ **Maleru Petru**, ce activează în baza Statutului, numită în continuare **„Angajator”**, pe de o parte, și
Dl. **{{NUME_COMPLET}}**, cetățenie **{{CETATENIA}}**, născut la data de **{{DATA_NASTERII}}**, locul nașterii **{{LOCUL_NASTERII}}**, seria/număr document de călătorie **{{NR_PASAPORT}}**, valabil până la **{{VALABILITATEA}}**, numit în continuare **„Salariat”**, pe de altă parte,
conducându-se de prevederile art. 45–94 din Codul muncii (Legea nr. 154-XV din 28.03.2003) și de art. 43 (indice 1–17) din Legea nr. 200 din 16.07.2010 privind regimul străinilor în Republica Moldova, au încheiat prezentul Contract individual de muncă, convenind asupra următoarelor:""",
"""**1. Contracting Parties**
“Europersonal” SRL, headquartered in Chisinau Municipality, Mitropolit Gavriil Bănulescu-Bodoni Street, No. 43a, Apt. 14, VAT ID 1016600043760, represented by the administrative director Maleru Petru, acting under the Articles of Association, hereinafter referred to as the “Employer”, on one side, and
Mr. **{{NUME_COMPLET}}**, citizenship **{{CETATENIA}}, born on {{DATA_NASTERII}},** place of birth **{{LOCUL_NASTERII}}, travel document series/number {{NR_PASAPORT}}**, valid until **{{VALABILITATEA}}** hereinafter referred to as the “Employee”, on the other side,
acting in accordance with the provisions of Articles 45–94 of the Labour Code (Law No. 154-XV of 28.03.2003) and Article 43 (indices 1–17) of Law No. 200 of 16.07.2010 regarding the legal regime of foreigners in the Republic of Moldova, have concluded this Individual Labor Contract, agreeing on the following:"""))
S.append((
"""**2. Obiectul contractului. Funcția**
2.1. Salariatul este angajat în funcția de **{{FUNCTIA_RO}}.**
2.2. Atribuțiile de serviciu sunt cele prevăzute în **Fișa postului**, parte aplicabilă raportului de muncă.""",
"""**2. Subject of the Contract. Position**
2.1. The Employee is hired in the position of **{{FUNCTIA_EN}}**
2.2. The job duties are those set out in the Job Description, applicable to the employment relationship."""))
S.append((
"""**3. Locul de muncă. Clauza de mobilitate**
3.1. Părțile convin că activitatea Salariatului nu se va desfășura într-un loc fix, ci va implica deplasări în diverse locații stabilite de Angajator, în funcție de necesități.
3.2. Angajatorul are dreptul să stabilească și să modifice, cu respectarea prevederilor legale, locul concret de desfășurare a activității în aria geografică aferentă funcției și atribuțiilor.""",
"""**3. Workplace. Mobility Clause**
3.1. The parties agree that the Employee’s activity will not be carried out in a fixed place, but will involve travel to various locations established by the Employer, depending on needs.
3.2. The Employer has the right to establish and modify, in compliance with the legal provisions, the specific place where the activity is carried out within the geographical area related to the position and duties."""))
S.append((
"""**4. Tipul muncii**
4.1. Munca este **de bază**.""",
"""**4. Type of Work**
4.1. The work is **basic** employment."""))
S.append((
"""**5. Durata contractului**
5.1. Prezentul contract se încheie pe durată determinată de **1 (un) an**.""",
"""**5. Contract Term**
5.1. This contract is concluded for a fixed term of **1 (one) year.**"""))
S.append((
"""**6. Perioada de probă**
6.1. Perioada de probă este de **30 (treizeci) zile calendaristice**.
6.2. Pe durata probei, Salariatul beneficiază de toate drepturile/garanțiile prevăzute de lege și are aceleași obligații ca ceilalți salariați; perioada se include în vechime.
6.3. Evaluarea în probă se face de conducătorul nemijlocit, pe baza criteriilor funcției (calificare/abilități, calitate/volum, disciplină, integrare). La solicitare, rezultatele se consemnează în scris.
6.4. În caz de absență justificată (medical, concediu, suspendare etc.), proba se suspendă și se prelungește corespunzător.
6.5. Încetarea CIM în perioada de probă se face conform Codului muncii.""",
"""**6. Probation Period**
6.1. The probation period is **30 (thirty) calendar days**.
6.2. During probation, the Employee enjoys all rights/guarantees provided by law and has the same obligations as the other employees; the period is included in length of service.
6.3. The probation assessment is carried out by the direct supervisor, based on the position criteria (qualification/skills, quality/volume, discipline, integration). Upon request, the results are recorded in writing.
6.4. In the event of justified absence (medical leave, annual leave, suspension, etc.), the probation is suspended and extended accordingly.
6.5. Termination of the individual labor contract during the probation period is carried out in accordance with the Labour Code."""))
S.append((
"""**7. Sarcini de bază**
Salariatul are, fără a se limita la, următoarele sarcini:
a) organizarea rațională a activității;
b) respectarea disciplinei de muncă;
c) păstrarea curățeniei și ordinii la locul de muncă;
d) sporirea competenței profesionale;
e) îndeplinirea atribuțiilor în strictă conformitate cu legea și regulamentele interne;
f) păstrarea confidențialității datelor/informațiilor accesate în exercitarea funcției.""",
"""**7. Main Tasks**
The Employee has, without limitation, the following tasks:
a) rational organization of activity;
b) compliance with work discipline;
c) maintaining cleanliness and order at the workplace;
d) improving professional competence;
e) performance of duties in strict compliance with the law and internal regulations;
f) maintaining the confidentiality of the data/information accessed in the performance of the position."""))
S.append((
"""**8. Drepturile Salariatului**
Salariatul are drepturile prevăzute de legislația muncii, inclusiv:
a) încheierea, modificarea, suspendarea și încetarea CIM în condițiile legii;
b) muncă conform contractului;
c) condiții de muncă sigure și igienice;
d) salariu conform calificării și muncii efectuate;
e) odihnă (program normal, repaus, sărbători, concedii plătite);
f) informare privind SSM;
g) adresare către Angajator și organele competente;
h) formare profesională;
i) apărarea drepturilor prin metode permise;
j) soluționarea litigiilor și dreptul la grevă conform legii;
k) repararea prejudiciului material și moral conform legii.""",
"""**8. Employee Rights**
The Employee has the rights provided by labor legislation, including:
a) conclusion, amendment, suspension, and termination of the individual labor contract under the conditions of the law;
b) work in accordance with the contract;
c) safe and hygienic working conditions;
d) salary according to qualification and work performed;
e) rest (normal working schedule, weekly rest, public holidays, paid leave);
f) information regarding occupational health and safety;
g) addressing the Employer and the competent authorities;
h) professional training;
i) protection of rights by permitted methods;
j) resolution of disputes and the right to strike in accordance with the law;
k) compensation for material and moral damage in accordance with the law."""))
S.append((
"""**9. Obligațiile Salariatului**
a) îndeplinirea conștiincioasă a obligațiilor;
b) respectarea normelor de muncă;
c) respectarea disciplinei;
d) respectarea cerințelor de protecție și igienă a muncii;
e) atitudine gospodărească față de bunurile Angajatorului și ale altor salariați;
f) informarea imediată despre orice situație periculoasă pentru viață/sănătate/patrimoniu.""",
"""**9. Employee Obligations**
a) conscientious performance of obligations;
b) compliance with work norms;
c) compliance with discipline;
d) compliance with occupational protection and workplace hygiene requirements;
e) prudent care for the Employer’s property and that of other employees;
f) immediate notification of any situation dangerous to life/health/property."""))
S.append((
"""**10. Drepturile Angajatorului**
a) încheiere/modificare/suspendare/încetare CIM conform legii;
b) solicitarea îndeplinirii obligațiilor conform fișei postului;
c) stimularea salariaților pentru muncă eficientă;
d) aplicarea răspunderii disciplinare conform legii.""",
"""**10. Employer Rights**
a) conclusion/amendment/suspension/termination of the individual labor contract in accordance with the law;
b) requiring the performance of obligations according to the job description;
c) encouraging employees for efficient work;
d) applying disciplinary liability in accordance with the law."""))
S.append((
"""**11. Obligațiile Angajatorului**
a) respectarea legii și a actelor normative aplicabile;
b) respectarea clauzelor CIM;
c) asigurarea muncii convenite;
d) asigurarea condițiilor corespunzătoare SSM;
e) asigurarea utilajului/instrumentelor necesare;
f) plată egală pentru muncă de valoare egală;
g) plata integrală a salariului la termenele stabilite;
h) examinarea sesizărilor și remedierea încălcărilor;
i) asigurare socială obligatorie;
j) repararea prejudiciilor conform legii.""",
"""**11. Employer Obligations**
a) compliance with the law and applicable normative acts;
b) compliance with the clauses of the individual labor contract;
c) ensuring the agreed work;
d) ensuring appropriate occupational health and safety conditions;
e) providing the necessary equipment/tools;
f) equal pay for work of equal value;
g) full payment of salary within the established deadlines;
h) examination of complaints and remedy of violations;
i) compulsory social insurance;
j) compensation for damages in accordance with the law."""))
S.append((
"""**12. Salariul. Modalitatea și frecvența plății. Ore suplimentare. Zile nelucrătoare**
12.1. Salariul lunar brut este de **{{SALARIU_RO}}**.
12.2. Plata salariului se face în MDL, prin **transfer bancar**, cel puțin **o dată pe lună**, în **ultimele zile lucrătoare** ale lunii pentru care se achită (sau conform graficului intern aprobat).
12.3. Munca suplimentară se compensează conform legislației și regulilor interne:
a) primele 2 ore – cel puțin **1,5x**;
b) orele următoare – cel puțin **2x**.
12.4. Munca în zile de repaus săptămânal și/sau sărbători nelucrătoare se compensează conform legii prin:
a) zi liberă plătită, la cererea Salariatului; sau
b) plată de cel puțin **dublu**, dacă nu se poate acorda zi liberă.""",
"""**12. Salary. Method and Frequency of Payment. Overtime. Non-working Days**
12.1. The gross monthly salary is **{{SALARIU_EN}}.**
12.2. Salary payment is made in MDL, by **bank transfer**, **at least once per month, during the last working days** of the month for which payment is made (or according to the approved internal schedule).
12.3. Overtime work is compensated in accordance with legislation and internal rules:
a) the first 2 hours – at least **1.5x**;
b) the following hours – at least **2x**.
12.4. Work performed on weekly rest days and/or public holidays is compensated in accordance with the law by:
a) a paid day off, at the Employee’s request; or
b) payment of at least **double**, if a day off cannot be granted."""))
S.append((
"""**13. Regimul de muncă și evidența timpului de muncă**
13.1. Programul de lucru:
● Luni–Vineri: 09:00–17:00
● Sâmbătă: 09:00–14:00
● Duminică: liber
13.2. Durata timpului de muncă: **40 ore/săptămână**.
13.3. Pauza de masă: **1 oră**, conform regulilor interne.
13.4. Evidența timpului de muncă se ține de Angajator conform procedurilor interne.""",
"""**13. Working Time Regime and Timekeeping**
13.1. Working schedule:
• Monday–Friday: 09:00–17:00
• Saturday: 09:00–14:00
• Sunday: day off
13.2. Working time duration: **40 hours/week.**
13.3. Lunch break: **1 hour**, according to internal rules.
13.4. Working time records are kept by the Employer according to internal procedures."""))
S.append((
"""**14. Repaus săptămânal**
14.1. Repausul săptămânal se acordă **duminica**.""",
"""**14. Weekly Rest**
14.1. Weekly rest is granted on Sunday."""))
S.append((
"""**15. Concedii. Concediu anual neutilizat/acumulat**
15.1. Concediul anual plătit: minimum **28 zile calendaristice**, cu excepția sărbătorilor nelucrătoare.
15.2. Concediul se acordă anual conform legii; Angajatorul nu poate refuza acordarea concediului anual **doi ani consecutivi**.
15.3. Zilele de concediu neutilizate pot fi programate și acordate ulterior conform legii și procedurilor interne, în termenul legal aplicabil.
15.4. La încetarea CIM, compensația pentru concediul neutilizat se achită conform legii.""",
"""**15. Leave. Unused/Accrued Annual Leave**
15.1. Annual paid leave: minimum **28 calendar days**, excluding public holidays.
15.2. Leave is granted annually according to the law; the Employer may not refuse granting annual leave f**or two consecutive years.**
15.3. Unused leave days may be scheduled and granted later according to the law and internal procedures, within the applicable legal term.
15.4. Upon termination of the individual labor contract, compensation for unused leave is paid according to the law."""))
S.append((
"""**16. Asigurări**
16.1. Asigurarea socială se efectuează conform legislației în vigoare.
16.2. Asigurarea medicală se efectuează conform legislației în vigoare.""",
"""**16. Insurance**
16.1. Social insurance is provided in accordance with the legislation in force.
16.2. Medical insurance is provided in accordance with the legislation in force."""))
S.append((
"""**17. Formare profesională / dezvoltare**
17.1. Angajatorul poate organiza/finanța programe de formare (cursuri interne/externe, SSM, instruiri specifice).
17.2. Condițiile concrete (durată, loc, tematică, costuri, obligații ulterioare – dacă e cazul) se stabilesc prin ordin și/sau act separat, în limitele legii.""",
"""**17. Professional Training / Development**
17.1. The Employer may organize/finance training programs (internal/external courses, occupational health and safety, job-specific training).
17.2. The specific conditions (duration, place, topics, costs, subsequent obligations – if applicable) are established by order and/or a separate act, within the limits of the law."""))
S.append((
"""**18. Nediscriminare. Prevenirea violenței și hărțuirii**
18.1. Se aplică principiul tratamentului egal și interdicția discriminării.
18.2. Salariatul respectă politicile/procedurile interne anti-hărțuire/anti-violență.
18.3. Angajatorul asigură condiții sigure și mecanisme interne de prevenire și raportare.
18.4. Acces egal și remunerare egală pentru muncă de valoare egală, indiferent de sex.
18.5. Încetarea CIM nu poate avea loc pe motive discriminatorii (sex, sarcină, concedii de maternitate/paternitate/îngrijire copil etc.).
18.6. Raportarea se poate face către HR/persoana desemnată, cu confidențialitate; raportarea de bună-credință nu atrage sancțiuni.""",
"""**18. Non-discrimination. Prevention of Violence and Harassment**
18.1. The principle of equal treatment and the prohibition of discrimination apply.
18.2. The Employee complies with the internal anti-harassment/anti-violence policies/procedures.
18.3. The Employer ensures safe conditions and internal prevention and reporting mechanisms.
18.4. Equal access and equal remuneration for work of equal value, regardless of sex.
18.5. Termination of the individual labor contract may not take place on discriminatory grounds (sex, pregnancy, maternity/paternity/childcare leave, etc.).
18.6. Reporting may be made to HR/the designated person, with confidentiality; good-faith reporting does not entail sanctions."""))
S.append((
"""**19. Protecția datelor cu caracter personal**
19.1. Angajatorul prelucrează date personale necesare raportului de muncă și obligațiilor legale.
19.2. Salariatul se obligă: (a) să respecte politicile; (b) să păstreze confidențialitatea inclusiv după încetare; (c) să nu divulge către terți; (d) să aplice măsuri de securitate.
19.3. Incidentele se raportează imediat.
19.4. Angajatorul aplică măsuri tehnice/organizatorice adecvate.
19.5. Prelucrarea categoriilor speciale de date se face doar în limitele legii/procedurilor interne.""",
"""**19. Personal Data Protection**
19.1. The Employer processes personal data necessary for the employment relationship and legal obligations.
19.2. The Employee undertakes: (a) to comply with policies; (b) to maintain confidentiality including after termination; (c) not to disclose to third parties; (d) to apply security measures.
19.3. Incidents are reported immediately.
19.4. The Employer applies appropriate technical/organizational measures.
19.5. Processing of special categories of data is carried out only within the limits of the law/internal procedures."""))
S.append((
"""**20. Securitate și sănătate în muncă (SSM)**
20.1. Salariatul respectă toate regulile SSM, participă la instruiri și utilizează echipamentele (dacă e cazul).
20.2. Instruirile și evidența pot fi realizate și electronic, conform procedurilor interne.
20.3. Salariatul raportează imediat accidente/incidente/riscuri.
20.4. Documentele SSM pot fi pe hârtie și/sau electronic.""",
"""**20. Occupational Health and Safety**
20.1. The Employee complies with all occupational health and safety rules, participates in training, and uses equipment (if applicable).
20.2. Trainings and records may also be carried out electronically, according to internal procedures.
20.3. The Employee immediately reports accidents/incidents/risks.
20.4. Occupational health and safety documents may be on paper and/or electronic."""))
S.append((
"""**21. Spații prietenoase pentru copii (dacă există)**
21.1. Angajatorul poate amenaja spații pentru copiii salariaților (peste 3 ani), pentru utilizare ocazională, conform regulilor interne.
21.2. Utilizarea nu constituie servicii autorizate; accesul este condiționat de reguli și intervale stabilite.""",
"""**21. Child-friendly Spaces (if available)**
21.1. The Employer may arrange spaces for employees’ children (over 3 years), for occasional use, according to internal rules.
21.2. Use does not constitute authorized services; access is conditional on rules and established time intervals."""))
S.append((
"""**22. Preaviz**
22.1. Demisie (inițiativa Salariatului): **14 zile calendaristice**, conform legii.
22.2. Încetare la inițiativa Angajatorului: preaviz conform legii, inclusiv:
a) **2 luni** la lichidare/reducere de personal;
b) **1 lună** la necorespundere din calificare insuficientă.
22.3. Alte situații – conform legislației.""",
"""**22. Notice**
22.1. Resignation (at the Employee’s initiative): 14 calendar days, according to the law.
22.2. Termination at the Employer’s initiative: notice according to the law, including:
a) 2 months in case of liquidation/staff reduction;
b) 1 month in case of unsuitability due to insufficient qualification.
22.3. Other situations – according to legislation."""))
S.append((
"""**23. Modificarea contractului**
23.1. Contractul poate fi modificat/completat numai prin **acord suplimentar** semnat de părți, care devine parte integrantă.
23.2. Se consideră modificare, fără a se limita la: durata, locul, specificul muncii, retribuirea, regimul de muncă/odihnă, funcția/calificarea, înlesnirile și modul de acordare.""",
"""**23. Amendment of the Contract**
23.1. The contract may be amended/supplemented only by a supplementary agreement signed by the parties, which becomes an integral part.
23.2. Amendments include, without limitation: term, workplace, nature of work, remuneration, working/rest schedule, position/qualification, benefits and the manner of granting them."""))
S.append((
"""**24. Modificare unilaterală (excepții)**
24.1. Modificarea unilaterală de către Angajator este posibilă doar în cazurile și condițiile prevăzute de Codul muncii, cu preavizul legal (de regulă, 2 luni, unde este aplicabil).""",
"""**24. Unilateral Amendment (exceptions)**
24.1. Unilateral amendment by the Employer is possible only in the cases and conditions provided by the Labour Code, with the legal notice (as a rule, 2 months, where applicable)."""))
S.append((
"""**25. Schimbarea temporară a locului/specificului muncii**
25.1. Locul de muncă poate fi schimbat temporar prin deplasare în interes de serviciu/detașare conform legii.
25.2. În cazurile prevăzute de art. 104 alin. (2) lit. a) și b) Codul muncii, Angajatorul poate schimba temporar locul și specificul muncii pentru max. 1 lună, conform legii.""",
"""**25. Temporary Change of Workplace/Nature of Work**
25.1. The workplace may be temporarily changed by service travel/secondment in accordance with the law.
25.2. In the cases provided by Article 104 paragraph (2) letters a) and b) of the Labour Code, the Employer may temporarily change the workplace and the nature of work for a maximum of 1 month, according to the law."""))
S.append((
"""**26. Suspendarea contractului**
26.1. Suspendarea poate surveni:
a) în circumstanțe independente de voința părților;
b) prin acordul părților;
c) la inițiativa uneia dintre părți, conform Codului muncii.""",
"""**26. Suspension of the Contract**
26.1. Suspension may occur:
a) in circumstances independent of the will of the parties;
b) by agreement of the parties;
c) at the initiative of one of the parties, in accordance with the Labour Code."""))
S.append((
"""**27. Încetarea contractului**
27.1. CIM poate înceta în condițiile Codului muncii (inclusiv circumstanțe independente / inițiativa unei părți).""",
"""**27. Termination of the Contract**
27.1. The individual labor contract may terminate under the conditions of the Labour Code (including circumstances independent / initiative of one party)."""))
S.append((
"""**28. Litigii**
28.1. Litigiile individuale de muncă se soluționează conform Codului muncii și altor acte normative.""",
"""**28. Disputes**
28.1. Individual labor disputes are resolved in accordance with the Labour Code and other normative acts."""))
S.append((
"""**29. Exemplare. Intrarea în vigoare**
29.1. Contractul este întocmit în **două exemplare** cu aceeași putere juridică: unul la Angajator, unul la Salariat.
29.2. Contractul intră în vigoare și își produce efectele **din momentul eliberării de către autoritățile competente ale RM a permisului de ședere în scop de muncă**.""",
"""**29. Copies. Entry into Force**
29.1. This contract is drawn up in **two originals** having the same legal force: one kept by the Employer, one by the Employee.
29.2. The contract enters into force and produces its effects **from the moment the competent authorities of the Republic of Moldova issue the residence permit for the purpose of work.**"""))

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(1.5)
sec.left_margin = sec.right_margin = Cm(1.5)
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = SIZE

table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
table.autofit = False

# rândul-titlu (centrat)
row = table.add_row()
fill_cell(row.cells[0], S[0][0], align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(row.cells[1], S[0][1], align=WD_ALIGN_PARAGRAPH.CENTER)

# secțiunile 1-29
for ro, en in S[1:]:
    row = table.add_row()
    fill_cell(row.cells[0], ro)
    fill_cell(row.cells[1], en)
    for c in row.cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# lățimi egale
for row in table.rows:
    row.cells[0].width = Cm(9.0)
    row.cells[1].width = Cm(9.0)

# blocul de semnături (rânduri cu celule unite, lățime întreagă)
def merged_row(text, align=WD_ALIGN_PARAGRAPH.LEFT):
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[1])
    fill_cell(cell, text, align=align)
    return cell

merged_row("**Date de identificare / Identification details & semnături / signatures**",
           align=WD_ALIGN_PARAGRAPH.CENTER)
merged_row("""**ANGAJATOR / EMPLOYER**
„Europersonal” SRL
Adresa/Address: Chișinău, str. Mitropolit Gavriil Bănulescu-Bodoni, nr. 43a, ap. 14
ID/VAT: 1016600043760
Funcție/Title: Director administrativ / Administrative Director
Nume/Name: **Maleru Petru**
Semnătura/Signature: ____________________  L.Ș./Stamp: _________

**SALARIAT / EMPLOYEE**
Nume/Name: **{{NUME_COMPLET}}**
Semnătura/Signature: ____________________""")
merged_row("""**Confirmare luare la cunoștință / Acknowledgment**
**RO:** Am luat cunoştinţă cu regulamentul intern, cerinţele privind securitatea şi sănătatea în muncă, securitatea anti-incendiu, precum şi cu regulile necesare pentru îndeplinirea obligaţiunilor mele de muncă.
**EN:** I acknowledge the internal regulations, occupational health and safety requirements, fire safety rules, and the rules necessary for performing my job duties.

Semnătura Salariatului / Employee signature: ____________________""")

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates", "CIM.docx")
doc.save(out)
print("Șablon salvat:", out)
