"""Verificări de consistență ale pachetului de documente.

Regula centrală (R1): adresa din Contractul de Comodat trebuie să apară
IDENTIC în toate documentele care conțin adresa de cazare.
"""
import io
import re
import unicodedata

import docx


def _norm(s: str) -> str:
    s = unicodedata.normalize("NFC", s)
    s = (s.replace("ş", "ș").replace("ţ", "ț")
          .replace("Ş", "Ș").replace("Ţ", "Ț"))
    return re.sub(r"\s+", " ", s).strip().lower().rstrip(".")


def _doc_text(docx_bytes: bytes) -> str:
    d = docx.Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


DOCS_REQUIRING_ADDRESS = ("CONFIRMARE", "ANEXA", "MODEL_M")
DOCS_REQUIRING_PASSPORT = ("DEMERS", "ANEXA", "MODEL_M", "INVITATIE", "CIM")


def check_package_consistency(files: dict, package) -> list[str]:
    warnings = []
    expected_addr = _norm(package.accommodation.address)
    expected_name = _norm(package.passport.full_name)
    expected_passport = package.passport.passport_number.upper()

    for fname, data in files.items():
        text = _doc_text(data)
        ntext = _norm(text)

        must_have_address = any(k in fname.upper() for k in DOCS_REQUIRING_ADDRESS)
        if must_have_address and expected_addr not in ntext:
            warnings.append(
                f"⚠️ {fname}: adresa de cazare NU corespunde cu cea din "
                f"Contractul de Comodat („{package.accommodation.address}”)."
            )

        # numele: fie complet, fie despărțit (Numele / Prenumele — INVITATIE, Model M)
        name_ok = (expected_name in ntext
                   or (_norm(package.passport.surname) in ntext
                       and _norm(package.passport.given_name) in ntext))
        if not name_ok:
            warnings.append(f"⚠️ {fname}: numele complet „{package.passport.full_name}” nu a fost găsit.")

        must_have_passport = any(k in fname.upper() for k in DOCS_REQUIRING_PASSPORT)
        compact = re.sub(r"[\s.]", "", text.upper())
        if (must_have_passport
                and expected_passport not in compact
                and package.passport.passport_no_digits not in compact):
            warnings.append(f"⚠️ {fname}: numărul de pașaport {expected_passport} nu a fost găsit.")

        other = {"NEPAL": ("BANGLADESH", "BANGLA"), "BANGLADESH": ("NEPAL", "NEPALI")}
        for wrong in other[package.passport.country.value]:
            if re.search(rf"\b{wrong}", text, re.I):
                warnings.append(
                    f"⚠️ {fname}: conține „{wrong}” — nepotrivit pentru un lucrător din "
                    f"{package.passport.country_ro} (eroare tip copy-paste)."
                )
    return warnings


def check_address_vs_contract(ui_address: str, contract_address) -> str | None:
    if contract_address and _norm(ui_address) != _norm(contract_address):
        return (f"⚠️ Adresa folosită („{ui_address}”) diferă de cea extrasă din "
                f"Contractul de Comodat („{contract_address}”). Documentele se "
                f"generează cu adresa introdusă — verificați!")
    return None
