"""
Creates Romanian-language police document templates for Voice Intake v0.0.1.
Run once: python template_builder.py

Each template uses {{field_name}} placeholders that doc_generator.py fills in.
Replace these .docx files with your official Moldovan signed templates when ready.

Templates (Declarații & Sesizări):
  1. plangere.docx         — Plângere Penală (victim complaint)
  2. declaratie.docx       — Declarație de Victimă (victim statement)
  3. proces_verbal.docx    — Proces-Verbal de Sesizare (officer record)
  4. nota_explicativa.docx — Notă Explicativă (person of interest)
  5. sesizare_penala.docx  — Sesizare Penală (complaint to prosecutor)

Templates (Rapoarte GO):
  6. raport_go.docx        — Raport Grupă Operativă / Apel 112
  7. raport_pierdere.docx  — Raport privind pierderea materialului
  8. raport_droguri.docx   — Raport predare droguri / substanțe
  9. raport_transmitere.docx — Raport transmitere material alt sector
 10. raport_cautare.docx   — Raport anunțare în căutare

Templates (Note & Interpelări):
 11. nota_informativa.docx — Notă informativă către șef
 12. interpelare.docx      — Interpelare către instituție
 13. prelungire.docx       — Raport de prelungire termen examinare

Templates (Decizii & Încheieri):
 14. decizie_clasare.docx  — Decizie clasare / relații civile
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import TEMPLATES_DIR


def _heading(doc, text, size=13, bold=True, center=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def _margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)


def _field_row(table, label, placeholder):
    row = table.add_row()
    row.cells[0].text = label + ":"
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = f"{{{{{placeholder}}}}}"
    return row


def _table2col(doc):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(10)
    return tbl


# ─── 1. Plângere Penală ────────────────────────────────────────────────────────

def build_plangere():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLIȚIE", 12)
    doc.add_paragraph()
    _heading(doc, "PLÂNGERE PENALĂ", 14)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. dosar: {{numar_dosar}}\n")
    p.add_run("Data: {{data_intocmirii}}    Ora: {{ora_intocmirii}}")

    doc.add_paragraph()
    _heading(doc, "I. DATE PRIVIND VICTIMA / PETENTUL", 11, center=False)
    tbl = _table2col(doc)
    for label, ph in [
        ("Nume",              "victima_nume"),
        ("Prenume",           "victima_prenume"),
        ("IDNP / Buletin",    "victima_idnp"),
        ("Data nașterii",     "victima_data_nastere"),
        ("Adresa domiciliu",  "victima_adresa"),
        ("Telefon",           "victima_telefon"),
    ]:
        _field_row(tbl, label, ph)

    doc.add_paragraph()
    _heading(doc, "II. DESCRIEREA INCIDENTULUI", 11, center=False)
    tbl2 = _table2col(doc)
    for label, ph in [
        ("Tip infracțiune",    "tip_incident"),
        ("Data incidentului",  "data_incident"),
        ("Ora incidentului",   "ora_incident"),
        ("Locul incidentului", "locul_incident"),
        ("Detalii loc",        "locul_detalii"),
    ]:
        _field_row(tbl2, label, ph)

    doc.add_paragraph()
    _heading(doc, "III. DESCRIEREA EVENIMENTULUI", 11, center=False)
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    _heading(doc, "IV. DESCRIEREA FĂPTUITORULUI", 11, center=False)
    doc.add_paragraph("{{faptuitor_descriere}}")
    doc.add_paragraph("{{faptuitor_detalii}}")

    doc.add_paragraph()
    _heading(doc, "V. MARTORI", 11, center=False)
    doc.add_paragraph("{{martori}}")

    doc.add_paragraph()
    _heading(doc, "VI. BUNURI SUSTRASE / PREJUDICIU", 11, center=False)
    tbl3 = _table2col(doc)
    _field_row(tbl3, "Bunuri sustrase",  "bunuri_sustrase")
    _field_row(tbl3, "Valoare estimată", "valoare_estimata")

    doc.add_paragraph()
    _heading(doc, "VII. SOLICITĂRI", 11, center=False)
    doc.add_paragraph("{{masuri_solicitate}}")

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Subsemnatul/a declar că cele menționate sunt adevărate și îmi asum răspunderea "
              "conform legislației în vigoare.\n\n")
    p.add_run("Semnătura petentului:    ________________________\n")
    p.add_run("Ofițer de serviciu:       ________________________")

    path = os.path.join(TEMPLATES_DIR, "plangere.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 2. Declarație de Victimă ─────────────────────────────────────────────────

def build_declaratie():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLIȚIE", 12)
    doc.add_paragraph()
    _heading(doc, "DECLARAȚIE", 14)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Data: {{data_intocmirii}}    Ora: {{ora_intocmirii}}")

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("Eu, subsemnatul/a ")
    r = intro.add_run("{{victima_prenume}} {{victima_nume}}")
    r.bold = True
    intro.add_run(
        ", născut/ă la {{victima_data_nastere}}, domiciliat/ă la adresa {{victima_adresa}}, "
        "posesor/posesoare al/a buletinului de identitate {{victima_idnp}}, "
        "telefon {{victima_telefon}}, declar următoarele:"
    )

    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    doc.add_paragraph("{{alte_detalii}}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Cele relatate mai sus corespund adevărului. Am fost informat/ă că declarațiile "
        "false atrag răspundere penală conform legislației Republicii Moldova.\n\n"
    )
    p.add_run("Semnătura declarantului:                  ________________________\n")
    p.add_run("Ofițer care a recepționat declarația:  ________________________")

    path = os.path.join(TEMPLATES_DIR, "declaratie.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 3. Proces-Verbal de Sesizare ─────────────────────────────────────────────

def build_proces_verbal():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLIȚIE", 12)
    doc.add_paragraph()
    _heading(doc, "PROCES-VERBAL DE SESIZARE", 14)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. dosar: {{numar_dosar}}\nData: {{data_intocmirii}}    Ora: {{ora_intocmirii}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofițerul de serviciu al Inspectoratului de Poliție a recepționat sesizarea cu privire la "
        "{{tip_incident}}, produs la data de {{data_incident}}, ora {{ora_incident}}, "
        "la locul {{locul_incident}}."
    )

    doc.add_paragraph()
    _heading(doc, "PERSOANA VĂTĂMATĂ:", 11, center=False)
    doc.add_paragraph(
        "Nume/Prenume: {{victima_prenume}} {{victima_nume}}\n"
        "IDNP: {{victima_idnp}}\nTelefon: {{victima_telefon}}"
    )

    doc.add_paragraph()
    _heading(doc, "DESCRIEREA FAPTEI:", 11, center=False)
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    _heading(doc, "PREJUDICIU:", 11, center=False)
    doc.add_paragraph("Bunuri sustrase: {{bunuri_sustrase}}\nValoare estimată: {{valoare_estimata}} MDL")

    doc.add_paragraph()
    _heading(doc, "MĂSURI LUATE / SOLICITATE:", 11, center=False)
    doc.add_paragraph("{{masuri_solicitate}}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Ofițer de serviciu:  ________________________\nGrad, Nume: ________________________")

    path = os.path.join(TEMPLATES_DIR, "proces_verbal.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 4. Notă Explicativă ──────────────────────────────────────────────────────

def build_nota_explicativa():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLIȚIE", 12)
    doc.add_paragraph()
    _heading(doc, "NOTĂ EXPLICATIVĂ", 14)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Data: {{data_intocmirii}}    Ora: {{ora_intocmirii}}")

    doc.add_paragraph()
    _heading(doc, "I. DATE PRIVIND PERSOANA AUDIATĂ", 11, center=False)
    tbl = _table2col(doc)
    for label, ph in [
        ("Nume",             "victima_nume"),
        ("Prenume",          "victima_prenume"),
        ("IDNP / Buletin",   "victima_idnp"),
        ("Data nașterii",    "victima_data_nastere"),
        ("Adresa domiciliu", "victima_adresa"),
        ("Telefon",          "victima_telefon"),
    ]:
        _field_row(tbl, label, ph)

    doc.add_paragraph()
    _heading(doc, "II. FAPTA PENTRU CARE ESTE AUDIATĂ PERSOANA", 11, center=False)
    tbl2 = _table2col(doc)
    for label, ph in [
        ("Tip faptă",   "tip_incident"),
        ("Data faptei", "data_incident"),
        ("Ora faptei",  "ora_incident"),
        ("Locul faptei","locul_incident"),
    ]:
        _field_row(tbl2, label, ph)

    doc.add_paragraph()
    _heading(doc, "III. EXPLICAȚIILE PERSOANEI AUDIATE", 11, center=False)
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    _heading(doc, "IV. ALTE DETALII MENȚIONATE", 11, center=False)
    doc.add_paragraph("{{alte_detalii}}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Subsemnatul/a confirm că explicațiile de mai sus îmi aparțin și corespund adevărului. "
        "Am fost informat/ă că declarațiile false atrag răspundere penală.\n\n"
    )
    p.add_run("Semnătura persoanei audiate:  ________________________\n")
    p.add_run("Ofițer auditor:                        ________________________")

    path = os.path.join(TEMPLATES_DIR, "nota_explicativa.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 5. Sesizare Penală ────────────────────────────────────────────────────────

def build_sesizare_penala():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLIȚIE", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. dosar: {{numar_dosar}}\nData: {{data_intocmirii}}")

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p2.add_run("Către: PROCURATURA RAIONULUI / MUNICIPIULUI")
    r.bold = True

    doc.add_paragraph()
    _heading(doc, "SESIZARE PENALĂ", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "Inspectoratul de Poliție, în conformitate cu prevederile art. 262-263 din Codul de "
        "Procedură Penală al Republicii Moldova, sesizează organul de urmărire penală cu privire "
        "la infracțiunea de {{tip_incident}}, comisă la data de {{data_incident}}, ora "
        "{{ora_incident}}, la locul {{locul_incident}}."
    )

    doc.add_paragraph()
    _heading(doc, "I. PERSOANA VĂTĂMATĂ", 11, center=False)
    tbl = _table2col(doc)
    for label, ph in [
        ("Nume / Prenume",   "victima_prenume"),
        ("IDNP",             "victima_idnp"),
        ("Adresa",           "victima_adresa"),
        ("Telefon",          "victima_telefon"),
    ]:
        _field_row(tbl, label, ph)

    doc.add_paragraph()
    _heading(doc, "II. DESCRIEREA FAPTEI ȘI A CIRCUMSTANȚELOR", 11, center=False)
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    _heading(doc, "III. PREJUDICIUL CAUZAT", 11, center=False)
    tbl2 = _table2col(doc)
    _field_row(tbl2, "Bunuri sustrase",  "bunuri_sustrase")
    _field_row(tbl2, "Valoare estimată", "valoare_estimata")

    doc.add_paragraph()
    _heading(doc, "IV. TEMEIUL LEGAL", 11, center=False)
    doc.add_paragraph("{{temeiul_legal}}")

    doc.add_paragraph()
    _heading(doc, "V. SOLICITARE", 11, center=False)
    doc.add_paragraph(
        "În temeiul celor expuse, solicităm pornirea urmăririi penale și luarea "
        "măsurilor prevăzute de lege.\n\n{{masuri_solicitate}}"
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Șef Inspectorat / Ofițer de serviciu:\n")
    p.add_run("Grad, Nume:       ________________________\n")
    p.add_run("Semnătura:         ________________________\n")
    p.add_run("Stampila unitatii: ________________________")

    path = os.path.join(TEMPLATES_DIR, "sesizare_penala.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 6. Raport Grupa Operativa / Apel 112 ───────────────────────────────────

def build_raport_go():
    doc = Document()
    _margins(doc)

    # Right-aligned addressee header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Domnului {{sef_grad}}\n").bold = True
    p.add_run("{{sef_prenume_nume}}\n").bold = True
    p.add_run("Sef al Inspectoratului de Politie Botanica")

    doc.add_paragraph()
    _heading(doc, "R A P O R T", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "Va raportez ca la data de {{data_incident}}, ora {{ora_incident}}, "
        "la adresa {{locul_incident}}, a avut loc {{tip_incident}}. "
        "A sesizat: {{victima_prenume}} {{victima_nume}}, tel. {{victima_telefon}}."
    )
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fapt despre ce Va raportez.")

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Ofiter de investigatii:\n")
    p2.add_run("{{ofiters_prenume}} {{ofiters_nume}}")

    path = os.path.join(TEMPLATES_DIR, "raport_go.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 7. Raport Pierdere ──────────────────────────────────────────────────────

def build_raport_pierdere():
    doc = Document()
    _margins(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Domnului {{sef_grad}}\n").bold = True
    p.add_run("{{sef_prenume_nume}}\n").bold = True
    p.add_run("Sef al Inspectoratului de Politie Botanica")

    doc.add_paragraph()
    _heading(doc, "R A P O R T", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "Va raportez ca la data de {{data_incident}}, ora {{ora_incident}}, "
        "la adresa {{locul_incident}}, s-a constatat pierderea / disparitia urmatoarelor bunuri: "
        "{{bunuri_sustrase}} (valoare estimata: {{valoare_estimata}} MDL). "
        "Persoana: {{victima_prenume}} {{victima_nume}}, IDNP {{victima_idnp}}, "
        "tel. {{victima_telefon}}."
    )
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fapt despre ce Va raportez.")

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Ofiter de investigatii:\n")
    p2.add_run("{{ofiters_prenume}} {{ofiters_nume}}")

    path = os.path.join(TEMPLATES_DIR, "raport_pierdere.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 8. Raport Droguri / Substante ──────────────────────────────────────────

def build_raport_droguri():
    doc = Document()
    _margins(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Domnului {{sef_grad}}\n").bold = True
    p.add_run("{{sef_prenume_nume}}\n").bold = True
    p.add_run("Sef al Inspectoratului de Politie Botanica")

    doc.add_paragraph()
    _heading(doc, "R A P O R T", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "Va raportez ca la data de {{data_incident}}, ora {{ora_incident}}, "
        "la adresa {{locul_incident}}, cetateanul/a {{victima_prenume}} {{victima_nume}}, "
        "n. {{victima_data_nastere}}, domiciliat/a la {{victima_adresa}}, IDNP {{victima_idnp}}, "
        "a predat de buna voie urmatoarele substante: {{substanta_descriere}}."
    )
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fapt despre ce Va raportez.")

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Ofiter de investigatii:\n")
    p2.add_run("{{ofiters_prenume}} {{ofiters_nume}}")

    path = os.path.join(TEMPLATES_DIR, "raport_droguri.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 9. Raport Transmitere Material Alt Sector ───────────────────────────────

def build_raport_transmitere():
    doc = Document()
    _margins(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Domnului {{sef_grad}}\n").bold = True
    p.add_run("{{sef_prenume_nume}}\n").bold = True
    p.add_run("Sef al Inspectoratului de Politie Botanica")

    doc.add_paragraph()
    _heading(doc, "R A P O R T", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "Va raportez ca materialul REI-2 nr. {{numar_dosar}} privind {{tip_incident}}, "
        "inregistrat la {{data_incident}}, referitor la cetateanul/a "
        "{{victima_prenume}} {{victima_nume}}, a fost transmis spre competenta solutionare "
        "la {{sector_destinatie}}, intrucit fapta s-a produs la {{locul_incident}}."
    )
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fapt despre ce Va raportez.")

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Ofiter de investigatii:\n")
    p2.add_run("{{ofiters_prenume}} {{ofiters_nume}}")

    path = os.path.join(TEMPLATES_DIR, "raport_transmitere.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 10. Raport Cautare Persoana ─────────────────────────────────────────────

def build_raport_cautare():
    doc = Document()
    _margins(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Domnului {{sef_grad}}\n").bold = True
    p.add_run("{{sef_prenume_nume}}\n").bold = True
    p.add_run("Sef al Inspectoratului de Politie Botanica")

    doc.add_paragraph()
    _heading(doc, "R A P O R T", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "Va raportez ca la data de {{data_incident}}, ora {{ora_incident}}, "
        "cetateanul/a {{victima_prenume}} {{victima_nume}}, n. {{victima_data_nastere}}, "
        "domiciliat/a la {{victima_adresa}}, tel. {{victima_telefon}}, "
        "a disparut de la adresa {{locul_incident}}. "
        "Descriere fizica: {{faptuitor_descriere}}."
    )
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Rog anuntarea in cautare conform procedurii stabilite.")

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Ofiter de investigatii:\n")
    p2.add_run("{{ofiters_prenume}} {{ofiters_nume}}")

    path = os.path.join(TEMPLATES_DIR, "raport_cautare.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 11. Nota Informativa ────────────────────────────────────────────────────

def build_nota_informativa():
    doc = Document()
    _margins(doc)

    _heading(doc, "NOTA INFORMATIVA!", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "Domnule Sef sa traiti,\n\n"
        "Va aduc la cunostinta ca la data de {{data_incident}}, ora {{ora_incident}}, "
        "la {{locul_incident}}, sursa {{sursa_informatiei}} a semnalat {{tip_incident}}, "
        "implicind cetateanul/a {{victima_prenume}} {{victima_nume}}."
    )
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Cu respect,\n\n")
    p.add_run("{{ofiters_prenume}} {{ofiters_nume}}")

    path = os.path.join(TEMPLATES_DIR, "nota_informativa.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 12. Interpelare ─────────────────────────────────────────────────────────

def build_interpelare():
    doc = Document()
    _margins(doc)

    # Official letterhead table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    left = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]

    left.text = "REPUBLICA MOLDOVA\nMINISTERUL AFACERILOR INTERNE\nINSPECTORATUL GENERAL AL POLITIEI\nInspectoratul de Politie Botanica"
    for run in left.paragraphs[0].runs:
        run.bold = True

    right.text = "Nr. {{numar_dosar}}\ndin {{data_incident}}"
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{institutie_tip}}\n")
    p.add_run("{{institutie_nume}}\n")
    p.add_run("{{institutie_adresa}}")

    doc.add_paragraph()
    _heading(doc, "I N T E R P E L A R E", 14)
    doc.add_paragraph()

    doc.add_paragraph(
        "In cadrul examinarii materialului REI-2 nr. {{numar_dosar}} privind {{tip_incident}}, "
        "produs la data de {{data_incident}}, ora {{ora_incident}}, la {{locul_incident}}, "
        "va solicitam urmatoarele date / documente:"
    )
    doc.add_paragraph()
    doc.add_paragraph("{{date_solicitate}}")
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Termenul de raspuns — 10 zile lucratoare de la data receptiei prezentei."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Ofiter executor:\n")
    p2.add_run("{{ofiters_prenume}} {{ofiters_nume}}\n")
    p2.add_run("Tel: {{ofiters_telefon}}")

    path = os.path.join(TEMPLATES_DIR, "interpelare.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 13. Raport de Prelungire ─────────────────────────────────────────────────

def build_prelungire():
    doc = Document()
    _margins(doc)

    # Approval stamp at top right
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("APROB\n").bold = True
    p.add_run("{{sef_prenume_nume}}\n")
    p.add_run("Sef al Inspectoratului de Politie Botanica\n")
    p.add_run("___________________ (semnatura)\n")
    p.add_run("Data: {{data_incident}}")

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("Sefului Inspectoratului de Politie Botanica")

    doc.add_paragraph()
    _heading(doc, "R A P O R T", 14)
    doc.add_paragraph()

    _heading(doc, "A  C O N S T A T A T:", 12, center=False)
    doc.add_paragraph(
        "In evidenta Inspectoratului de Politie Botanica se afla inregistrat materialul "
        "REI-2 nr. {{numar_dosar}} din {{data_inregistrare}}, privind {{tip_incident}}, "
        "cu privire la cetateanul/a {{victima_prenume}} {{victima_nume}}."
    )

    doc.add_paragraph()
    _heading(doc, "S  O  L  I  C  I  T:", 12, center=False)
    doc.add_paragraph(
        "Prelungirea termenului de examinare a materialului REI-2 nr. {{numar_dosar}} "
        "in legatura cu:"
    )
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run("Coordonat:\n\n")
    p3.add_run("{{ofiters_prenume}} {{ofiters_nume}}\n")
    p3.add_run("Ofiter de investigatii")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.add_run(" P R E L U N G E S C \n")
    p4.add_run("pina la __________________\n")
    p4.add_run("{{sef_prenume_nume}}")

    path = os.path.join(TEMPLATES_DIR, "prelungire.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 14. Decizie Clasare / Relatii Civile ────────────────────────────────────

def build_decizie_clasare():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "D E C I Z I E", 16)
    doc.add_paragraph()

    _heading(doc, "A  C O N S T A T A T:", 12, center=False)
    doc.add_paragraph()

    doc.add_paragraph(
        "In evidenta Inspectoratului de Politie Botanica, sectia de investigatii, "
        "la {{data_inregistrare}} a fost inregistrat materialul REI-2 nr. {{numar_dosar}}, "
        "privind {{tip_incident}}, produs la data de {{data_incident}}, "
        "cu referire la cetateanul/a {{victima_prenume}} {{victima_nume}}, "
        "n. {{victima_data_nastere}}."
    )

    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    _heading(doc, "In baza celor expuse propun:", 12, center=False)
    doc.add_paragraph()

    p_c = doc.add_paragraph()
    p_c.add_run("1. {{concluzie_1}}")

    doc.add_paragraph()
    doc.add_paragraph()

    # Triple signature block
    tbl_sig = doc.add_table(rows=2, cols=3)
    tbl_sig.style = "Table Grid"
    headers = ["Ofiter de investigatii", "Sef sectie investigatii", "Sef Inspectorat"]
    for i, h in enumerate(headers):
        tbl_sig.rows[0].cells[i].text = h
        tbl_sig.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, val in enumerate(["{{ofiters_prenume}} {{ofiters_nume}}", "________________________", "________________________"]):
        tbl_sig.rows[1].cells[i].text = val

    path = os.path.join(TEMPLATES_DIR, "decizie_clasare.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 15. Incheiere ────────────────────────────────────────────────────────────

def build_incheiere():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    # Case number + date (right-aligned)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "Î N C H E I E R E", 16)
    doc.add_paragraph()

    # City/date line
    p_city = doc.add_paragraph()
    p_city.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_city.add_run("mun. Chișinău, {{data_incident}}")

    doc.add_paragraph()

    # Officer intro paragraph
    doc.add_paragraph(
        "Ofițer de investigații al Inspectoratului de Poliție Botanica, "
        "{{ofiters_prenume}} {{ofiters_nume}}, examinând materialele "
        "cauzei nr. {{numar_dosar}}, privind {{tip_incident}}, "
        "înregistrat la {{data_inregistrare}},"
    )

    doc.add_paragraph()
    _heading(doc, "A  C O N S T A T A T:", 12, center=False)
    doc.add_paragraph()

    # Facts paragraph
    doc.add_paragraph(
        "La {{data_inregistrare}} în evidența Inspectoratului de Poliție Botanica "
        "a fost înregistrat materialul privind {{tip_incident}}, cu referire la "
        "cetățeanul/a {{victima_prenume}} {{victima_nume}}, n. {{victima_data_nastere}}, "
        "domiciliat/ă la {{victima_adresa}}."
    )

    doc.add_paragraph()

    # Extended narrative (AI-filled)
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    _heading(doc, "Î N  T E M E I U L  C E L O R  E X P U S E:", 12, center=False)
    doc.add_paragraph()

    doc.add_paragraph(
        "În temeiul art. {{articol_legal}} din Codul de Procedură Penală al Republicii Moldova,"
    )

    doc.add_paragraph()
    _heading(doc, "D I S P U N:", 13, center=False)
    doc.add_paragraph()

    p_d = doc.add_paragraph()
    p_d.add_run("{{dispozitiv}}")

    doc.add_paragraph()

    # Legal article reference + appeal note
    p_appeal = doc.add_paragraph()
    p_appeal.add_run(
        "Prezenta încheiere poate fi contestată la procurorul ierarhic superior "
        "în termen de 10 zile din data emiterii."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Dual signature block
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.style = "Table Grid"
    tbl_sig.rows[0].cells[0].text = "Ofițer de investigații"
    tbl_sig.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[0].cells[1].text = "Coordonat: Șef secție investigații"
    tbl_sig.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[1].cells[0].text = "{{ofiters_prenume}} {{ofiters_nume}}"
    tbl_sig.rows[1].cells[1].text = "________________________"

    path = os.path.join(TEMPLATES_DIR, "incheiere.docx")
    doc.save(path)
    print(f"  [OK] {path}")



# ═══ v0.0.4 NEW TEMPLATES ════════════════════════════════════════════════════


# ─── 16. Proces-Verbal de Retinere ───────────────────────────────────────────

def build_pv_retinere():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "PROCES-VERBAL DE RETINERE", 14)
    _heading(doc, "(art. 167 CPP al Republicii Moldova)", 11, bold=False)
    doc.add_paragraph()

    p_city = doc.add_paragraph()
    p_city.add_run("mun. Chisinau, {{data_incident}},  ora {{ora_incident}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofiter de investigatii al Inspectoratului de Politie Botanica, "
        "{{ofiters_prenume}} {{ofiters_nume}}, in temeiul art. 167 CPP RM, "
        "am retinut in calitate de banuit pe cetateanul/a:"
    )

    doc.add_paragraph()
    tbl = _table2col(doc)
    _field_row(tbl, "Nume, Prenume", "victima_prenume}} {{victima_nume")
    _field_row(tbl, "Data nasterii", "victima_data_nastere")
    _field_row(tbl, "IDNP / Buletin", "victima_idnp")
    _field_row(tbl, "Domiciliu", "victima_adresa")
    _field_row(tbl, "Telefon", "victima_telefon")

    doc.add_paragraph()
    doc.add_paragraph("Temeiurile retinerii:")
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Banuitului i s-au adus la cunostinta drepturile conform art. 64 CPP RM, "
        "inclusiv dreptul de a tacea, dreptul la aparator, dreptul de a anunta "
        "rudele despre retinere. Starea fizica la retinere: {{stare_fizica}}."
    )

    doc.add_paragraph()
    doc.add_paragraph("Obiectii ale banuitului: {{obiectii}}")

    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.style = "Table Grid"
    tbl_sig.rows[0].cells[0].text = "Ofiter de investigatii"
    tbl_sig.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[0].cells[1].text = "Banuitului (semnatura)"
    tbl_sig.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[1].cells[0].text = "{{ofiters_prenume}} {{ofiters_nume}}"
    tbl_sig.rows[1].cells[1].text = "________________________"

    path = os.path.join(TEMPLATES_DIR, "pv_retinere.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 17. Proces-Verbal de Perchezitie ────────────────────────────────────────

def build_pv_perchezitie():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "PROCES-VERBAL DE PERCHEZITIE CORPORALA", 14)
    _heading(doc, "(art. 130 CPP al Republicii Moldova)", 11, bold=False)
    doc.add_paragraph()

    p_city = doc.add_paragraph()
    p_city.add_run("mun. Chisinau, {{data_incident}},  ora {{ora_incident}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofiter de investigatii {{ofiters_prenume}} {{ofiters_nume}}, "
        "in prezenta martorilor asistenti {{martori}}, "
        "in temeiul ordonantei din {{data_inregistrare}}, "
        "am efectuat perchezitia corporala a cetateanului/ei:"
    )

    doc.add_paragraph()
    tbl = _table2col(doc)
    _field_row(tbl, "Nume, Prenume", "victima_prenume}} {{victima_nume")
    _field_row(tbl, "IDNP", "victima_idnp")
    _field_row(tbl, "Locul perchezitiei", "locul_incident")

    doc.add_paragraph()
    doc.add_paragraph("Obiecte ridicate:")
    doc.add_paragraph("{{bunuri_sustrase}}")

    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Obiecte transmise spre examinare conform inventarului atasat. "
        "Obiectii: {{obiectii}}"
    )

    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=3)
    tbl_sig.style = "Table Grid"
    for i, h in enumerate(["Ofiter", "Martor asistent", "Persoana perchezitionata"]):
        tbl_sig.rows[0].cells[i].text = h
        tbl_sig.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, v in enumerate(["{{ofiters_prenume}} {{ofiters_nume}}", "________________________", "________________________"]):
        tbl_sig.rows[1].cells[i].text = v

    path = os.path.join(TEMPLATES_DIR, "pv_perchezitie.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 18. Proces-Verbal de Audiere ────────────────────────────────────────────

def build_pv_audiere():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "PROCES-VERBAL DE AUDIERE", 14)
    doc.add_paragraph()

    p_city = doc.add_paragraph()
    p_city.add_run("mun. Chisinau, {{data_incident}},  ora {{ora_incident}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofiter de investigatii {{ofiters_prenume}} {{ofiters_nume}} "
        "a audiat in calitate de {{calitate_persoana}} pe cetateanul/a:"
    )

    doc.add_paragraph()
    tbl = _table2col(doc)
    _field_row(tbl, "Nume, Prenume", "victima_prenume}} {{victima_nume")
    _field_row(tbl, "Data nasterii", "victima_data_nastere")
    _field_row(tbl, "IDNP", "victima_idnp")
    _field_row(tbl, "Domiciliu", "victima_adresa")
    _field_row(tbl, "Telefon", "victima_telefon")

    doc.add_paragraph()
    doc.add_paragraph(
        "Persoana audiata a fost avertizata ca declaratiile false atrag "
        "raspundere penala conform art. 312 CP RM."
    )

    doc.add_paragraph()
    doc.add_paragraph("DECLARATII:")
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")
    doc.add_paragraph()
    doc.add_paragraph("{{versiunea_proprie}}")

    doc.add_paragraph()
    doc.add_paragraph("Adaosuri / Obiectii: {{obiectii}}")

    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.style = "Table Grid"
    tbl_sig.rows[0].cells[0].text = "Ofiter de investigatii"
    tbl_sig.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[0].cells[1].text = "Persoana audiata"
    tbl_sig.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[1].cells[0].text = "{{ofiters_prenume}} {{ofiters_nume}}"
    tbl_sig.rows[1].cells[1].text = "________________________"

    path = os.path.join(TEMPLATES_DIR, "pv_audiere.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 19. Ordonanta de Refuz in Pornirea Urmaririi Penale ─────────────────────

def build_ordonanta_refuz():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "O R D O N A N T A", 16)
    _heading(doc, "de refuz in pornirea urmaririi penale", 12, bold=False)
    _heading(doc, "(art. 274 CPP al Republicii Moldova)", 11, bold=False)
    doc.add_paragraph()

    p_city = doc.add_paragraph()
    p_city.add_run("mun. Chisinau, {{data_incident}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofiter de investigatii al Inspectoratului de Politie Botanica, "
        "{{ofiters_prenume}} {{ofiters_nume}}, examinand sesizarea nr. {{numar_dosar}} "
        "inregistrata la {{data_inregistrare}}, din partea cetateanului/ei "
        "{{victima_prenume}} {{victima_nume}}, referitoare la {{tip_incident}},"
    )

    doc.add_paragraph()
    doc.add_paragraph("A  C O N S T A T A T:")
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "In temeiul art. 274 alin. (1) pct. {{articol_legal}} CPP RM, "
        "intrucat fapta nu intruneste elementele constitutive ale infractiunii / "
        "lipsesc temeiurile de pornire a urmaririi penale,"
    )

    doc.add_paragraph()
    doc.add_paragraph("D I S P U N:")
    doc.add_paragraph()
    doc.add_paragraph(
        "Refuzul in pornirea urmaririi penale pe sesizarea nr. {{numar_dosar}}.\n"
        "{{dispozitiv}}"
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Petentul {{victima_prenume}} {{victima_nume}} se informeaza despre "
        "prezenta ordonanta si dreptul de a o contesta procurorului ierarhic "
        "superior in termen de 10 zile."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.style = "Table Grid"
    tbl_sig.rows[0].cells[0].text = "Ofiter de investigatii"
    tbl_sig.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[0].cells[1].text = "Coordonat: Procuror"
    tbl_sig.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    tbl_sig.rows[1].cells[0].text = "{{ofiters_prenume}} {{ofiters_nume}}"
    tbl_sig.rows[1].cells[1].text = "________________________"

    path = os.path.join(TEMPLATES_DIR, "ordonanta_refuz.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 20. Raport de Serviciu ───────────────────────────────────────────────────

def build_raport_serviciu():
    doc = Document()
    _margins(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Domnului {{sef_grad}}\n{{sef_prenume_nume}}\nSef al Inspectoratului de Politie Botanica")
    run.bold = True

    doc.add_paragraph()
    _heading(doc, "R A P O R T   D E   S E R V I C I U", 14)
    doc.add_paragraph()

    p_date = doc.add_paragraph()
    p_date.add_run("Data: {{data_incident}}     Tura / Schimbul: {{ora_incident}}")

    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph("Evenimente deosebite:")
    doc.add_paragraph("{{alte_detalii}}")

    doc.add_paragraph()
    doc.add_paragraph("Masuri intreprinse:")
    doc.add_paragraph("{{masuri_solicitate}}")

    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Ofiter de serviciu:\n{{ofiters_prenume}} {{ofiters_nume}}")

    path = os.path.join(TEMPLATES_DIR, "raport_serviciu.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 21. Proces-Verbal de Constatare a Infractiunii ──────────────────────────

def build_pv_constatare():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "PROCES-VERBAL DE CONSTATARE A INFRACTIUNII", 14)
    _heading(doc, "(flagrant / la fata locului)", 11, bold=False)
    doc.add_paragraph()

    p_city = doc.add_paragraph()
    p_city.add_run("mun. Chisinau ({{locul_incident}}), {{data_incident}},  ora {{ora_incident}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofiter de investigatii {{ofiters_prenume}} {{ofiters_nume}}, "
        "deplasandu-se la locul faptei, a constatat savarsirea infractiunii de "
        "{{tip_incident}}, cu privire la cetateanul/a {{victima_prenume}} {{victima_nume}}."
    )

    doc.add_paragraph()
    doc.add_paragraph("Circumstantele constatate:")
    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph("Probe / corpuri delicte ridicate:")
    doc.add_paragraph("{{bunuri_sustrase}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Fapta se califica provizoriu conform art. {{articol_legal}} CP RM. "
        "Materialele urmeaza a fi inregistrate in SII in termenul legal."
    )

    doc.add_paragraph()
    doc.add_paragraph("Martori asistenti: {{martori}}")
    doc.add_paragraph("Obiectii: {{obiectii}}")

    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=3)
    tbl_sig.style = "Table Grid"
    for i, h in enumerate(["Ofiter de investigatii", "Martor asistent", "Martor asistent"]):
        tbl_sig.rows[0].cells[i].text = h
        tbl_sig.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, v in enumerate(["{{ofiters_prenume}} {{ofiters_nume}}", "________________________", "________________________"]):
        tbl_sig.rows[1].cells[i].text = v

    path = os.path.join(TEMPLATES_DIR, "pv_constatare.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 22. Proces-Verbal de Sechestru / Ridicare Bunuri ────────────────────────

def build_pv_sechestru():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "PROCES-VERBAL DE RIDICARE / SECHESTRU", 14)
    _heading(doc, "(art. 126-130 CPP al Republicii Moldova)", 11, bold=False)
    doc.add_paragraph()

    p_city = doc.add_paragraph()
    p_city.add_run("mun. Chisinau, {{data_incident}},  ora {{ora_incident}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofiter de investigatii {{ofiters_prenume}} {{ofiters_nume}}, "
        "in cadrul examinarii materialului nr. {{numar_dosar}} privind {{tip_incident}}, "
        "in prezenta martorilor asistenti {{martori}}, "
        "la adresa {{locul_incident}}, a ridicat / pus sub sechestru "
        "urmatoarele bunuri / documente apartinand cetateanului/ei "
        "{{victima_prenume}} {{victima_nume}}:"
    )

    doc.add_paragraph()
    doc.add_paragraph("Inventar bunuri ridicate:")
    doc.add_paragraph("{{bunuri_sustrase}}")
    doc.add_paragraph()
    doc.add_paragraph("Valoare estimata: {{valoare_estimata}} MDL")

    doc.add_paragraph()
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Bunurile ridicate au fost ambalate si sigilate. "
        "Obiectii ale persoanei: {{obiectii}}"
    )

    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=3)
    tbl_sig.style = "Table Grid"
    for i, h in enumerate(["Ofiter de investigatii", "Martor asistent", "Persoana de la care s-a ridicat"]):
        tbl_sig.rows[0].cells[i].text = h
        tbl_sig.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, v in enumerate(["{{ofiters_prenume}} {{ofiters_nume}}", "________________________", "________________________"]):
        tbl_sig.rows[1].cells[i].text = v

    path = os.path.join(TEMPLATES_DIR, "pv_sechestru.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── 23. Proces-Verbal de Examinare la Fata Locului ──────────────────────────

def build_pv_examinare():
    doc = Document()
    _margins(doc)

    _heading(doc, "REPUBLICA MOLDOVA", 11, bold=False)
    _heading(doc, "INSPECTORATUL DE POLITIE BOTANICA", 12)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Nr. {{numar_dosar}}\ndin {{data_inregistrare}}")

    doc.add_paragraph()
    _heading(doc, "PROCES-VERBAL DE EXAMINARE LA FATA LOCULUI", 14)
    _heading(doc, "(art. 118-121 CPP al Republicii Moldova)", 11, bold=False)
    doc.add_paragraph()

    p_city = doc.add_paragraph()
    p_city.add_run("{{locul_incident}}, {{data_incident}},  ora {{ora_incident}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Ofiter de investigatii {{ofiters_prenume}} {{ofiters_nume}}, "
        "in prezenta specialistului-criminalist si a martorilor asistenti {{martori}}, "
        "a examinat locul faptei privind {{tip_incident}}, "
        "in cauza cetateanului/ei {{victima_prenume}} {{victima_nume}}."
    )

    doc.add_paragraph()
    doc.add_paragraph("Descrierea locului examinat:")
    doc.add_paragraph("{{text_razvtat_ro}}")

    doc.add_paragraph()
    doc.add_paragraph("Probe / urme ridicate:")
    doc.add_paragraph("{{bunuri_sustrase}}")

    doc.add_paragraph()
    doc.add_paragraph("Fotografii / schite efectuate:")
    doc.add_paragraph("{{alte_detalii}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Examinarea s-a terminat la ora {{ora_incident}}. "
        "Obiectii: {{obiectii}}"
    )

    doc.add_paragraph()
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=3)
    tbl_sig.style = "Table Grid"
    for i, h in enumerate(["Ofiter de investigatii", "Martor asistent 1", "Martor asistent 2"]):
        tbl_sig.rows[0].cells[i].text = h
        tbl_sig.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, v in enumerate(["{{ofiters_prenume}} {{ofiters_nume}}", "________________________", "________________________"]):
        tbl_sig.rows[1].cells[i].text = v

    path = os.path.join(TEMPLATES_DIR, "pv_examinare.docx")
    doc.save(path)
    print(f"  [OK] {path}")


# ─── Runner ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Building Voice Intake v0.0.4 templates...")
    # v0.0.3 templates
    build_plangere()
    build_declaratie()
    build_proces_verbal()
    build_nota_explicativa()
    build_sesizare_penala()
    build_raport_go()
    build_raport_pierdere()
    build_raport_droguri()
    build_raport_transmitere()
    build_raport_cautare()
    build_nota_informativa()
    build_interpelare()
    build_prelungire()
    build_decizie_clasare()
    build_incheiere()
    # v0.0.4 new templates
    build_pv_retinere()
    build_pv_perchezitie()
    build_pv_audiere()
    build_ordonanta_refuz()
    build_raport_serviciu()
    build_pv_constatare()
    build_pv_sechestru()
    build_pv_examinare()
    print(f"\nAll 23 templates created in: {TEMPLATES_DIR}")
