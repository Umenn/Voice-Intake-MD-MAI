"""
file_extractor.py — extract plain text from uploaded files
Supported:
  document mode  — PDF (pdfplumber), DOCX, TXT/MD/CSV, scanned images (Tesseract)
  handwriting mode — images (JPG/PNG/WEBP) via GPT-4o Vision
v0.0.3
"""

import base64
import os
from pathlib import Path

# Load .env before reading any env vars — config.py may not have run yet at import time
try:
    from dotenv import load_dotenv
    load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env", override=False)
except ImportError:
    pass

IMAGE_EXTS    = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".tif"}
DOCUMENT_EXTS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv"}

# Point pytesseract at the Tesseract binary stored in .env (TESSERACT_CMD).
# This avoids adding Tesseract to the system PATH.
_tess_cmd = os.getenv("TESSERACT_CMD", "").strip()
if _tess_cmd:
    try:
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = _tess_cmd
    except ImportError:
        pass


def extract_text(file_path: str, mode: str = "document") -> dict:
    """
    Extract text from a file.
    mode: "document" — PDF/DOCX/TXT or Tesseract for images
          "handwriting" — GPT-4o Vision for images

    Returns {"filename", "text", "pages", "error"}
    """
    p = Path(file_path)
    suffix = p.suffix.lower()
    result = {"filename": p.name, "text": "", "pages": None, "error": None}

    try:
        if suffix in IMAGE_EXTS:
            if mode == "handwriting":
                result.update(_from_image_gpt(file_path))
            else:
                result.update(_from_image_tesseract(file_path))
        elif suffix == ".pdf":
            result.update(_from_pdf(file_path))
        elif suffix in (".docx", ".doc"):
            result.update(_from_docx(file_path))
        elif suffix in (".txt", ".md", ".csv"):
            result.update(_from_text(file_path))
        else:
            result["error"] = f"Format nesupport: {suffix}"
    except Exception as e:
        result["error"] = str(e)

    return result


# ── GPT-4o Vision (handwriting) ───────────────────────────────────────────────

def _from_image_gpt(path: str) -> dict:
    """Send image to GPT-4o vision and return extracted text."""
    import openai
    import config

    with open(path, "rb") as f:
        img_bytes = f.read()

    suffix = Path(path).suffix.lower().lstrip(".")
    mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png",
                "webp": "webp", "bmp": "png", "tiff": "png", "tif": "png"}
    mime = "image/" + mime_map.get(suffix, "jpeg")
    b64  = base64.b64encode(img_bytes).decode("utf-8")

    client = openai.OpenAI(api_key=config.OPENAI_API_KEY)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{b64}"},
                    },
                    {
                        "type": "text",
                        "text": (
                            "Esti un asistent pentru politia din Moldova. "
                            "Transcrie EXACT tot textul scris de mana din aceasta imagine. "
                            "Pastreaza structura originala (randuri noi, paragrafe). "
                            "Daca exista date (nume, date, adrese, numere de dosar), "
                            "subliniaza-le cu **bold**. "
                            "Raspunde DOAR cu textul transcris, fara explicatii suplimentare."
                        ),
                    },
                ],
            }
        ],
        max_tokens=2000,
    )

    text = response.choices[0].message.content or ""
    return {"text": text.strip(), "pages": 1}


# ── Tesseract OCR (scanned images) ───────────────────────────────────────────

def _from_image_tesseract(path: str) -> dict:
    """Extract text from a scanned image using Tesseract."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return {
            "text": "",
            "error": "pytesseract sau Pillow nu sunt instalate. Rulati: pip install pytesseract Pillow si instalati Tesseract de pe https://github.com/UB-Mannheim/tesseract/wiki",
        }

    try:
        img  = Image.open(path)
        # Try Romanian + Russian, fallback to eng
        text = pytesseract.image_to_string(img, lang="ron+rus", config="--psm 6")
    except pytesseract.TesseractNotFoundError:
        return {
            "text": "",
            "error": "Tesseract nu este instalat. Descarcati de la: https://github.com/UB-Mannheim/tesseract/wiki",
        }
    except Exception:
        # Fallback to English only
        try:
            text = pytesseract.image_to_string(img, config="--psm 6")
        except Exception as e:
            return {"text": "", "error": str(e)}

    return {"text": text.strip(), "pages": 1}


# ── PDF ───────────────────────────────────────────────────────────────────────

def _from_pdf(path: str) -> dict:
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                pages_text.append(t)
        combined = "\n\n".join(pages_text).strip()
        if combined:
            return {"text": combined, "pages": len(pages_text)}
        # PDF might be scanned — fall through to Tesseract
    except ImportError:
        pass

    # Scanned PDF: convert pages to images via Tesseract
    return _pdf_ocr_tesseract(path)


def _pdf_ocr_tesseract(path: str) -> dict:
    """OCR a scanned PDF page by page using Tesseract."""
    try:
        import pytesseract
        from PIL import Image
        import pdf2image  # optional — requires poppler
        pages = pdf2image.convert_from_path(path, dpi=200)
        texts = []
        for pg in pages:
            try:
                t = pytesseract.image_to_string(pg, lang="ron+rus", config="--psm 6")
            except Exception:
                t = pytesseract.image_to_string(pg, config="--psm 6")
            texts.append(t)
        return {"text": "\n\n".join(texts).strip(), "pages": len(texts)}
    except Exception:
        pass

    # Last resort: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        texts  = [p.extract_text() or "" for p in reader.pages]
        return {"text": "\n\n".join(texts).strip(), "pages": len(texts)}
    except ImportError:
        pass

    return {"text": "", "error": "Nu s-a putut extrage text din PDF (pdfplumber + Tesseract necesar)."}


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _from_docx(path: str) -> dict:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in paragraphs:
                    paragraphs.append(t)
    return {"text": "\n".join(paragraphs).strip(), "pages": None}


# ── Plain text ────────────────────────────────────────────────────────────────

def _from_text(path: str) -> dict:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return {"text": f.read().strip(), "pages": None}
