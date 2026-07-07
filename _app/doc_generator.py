"""
Fill a Word (.docx) template with extracted case data.
Templates use {{field_name}} placeholders in the document text.
"""

import os
import re
import copy
from datetime import datetime
from docx import Document
from docx.shared import Pt
from config import TEMPLATES_DIR, CASES_DIR


def _replace_in_paragraph(para, mapping: dict):
    """Replace all {{key}} occurrences in a paragraph, preserving formatting."""
    full_text = "".join(r.text for r in para.runs)
    replaced = _replace_placeholders(full_text, mapping)
    if replaced == full_text:
        return
    # Clear runs and put replaced text in first run
    for i, run in enumerate(para.runs):
        run.text = replaced if i == 0 else ""


def _replace_placeholders(text: str, mapping: dict) -> str:
    def replacer(m):
        key = m.group(1).strip()
        return mapping.get(key, m.group(0))
    return re.sub(r"\{\{(.+?)\}\}", replacer, text)


def _process_document(doc: Document, mapping: dict):
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, mapping)


def generate(template_name: str, data: dict, case_number: str = None) -> str:
    """
    Fill the given template with data and save to cases/ folder.

    Args:
        template_name: filename without extension (e.g. 'plangere')
        data:          dict with field values (from extractor.py)
        case_number:   optional case number string

    Returns:
        Path to the generated .docx file
    """
    template_path = os.path.join(TEMPLATES_DIR, f"{template_name}.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Build the full mapping (add meta-fields)
    ts = datetime.now()
    mapping = {**data}
    mapping.setdefault("data_intocmirii", ts.strftime("%d.%m.%Y"))
    mapping.setdefault("ora_intocmirii",  ts.strftime("%H:%M"))
    mapping.setdefault("numar_dosar",     case_number or "___________")

    _process_document(doc, mapping)

    # Save
    if case_number:
        fname = f"{case_number}_{template_name}_{ts.strftime('%Y%m%d_%H%M%S')}.docx"
    else:
        fname = f"{template_name}_{ts.strftime('%Y%m%d_%H%M%S')}.docx"

    out_path = os.path.join(CASES_DIR, fname)
    doc.save(out_path)
    return out_path


def list_templates() -> list[dict]:
    """Return list of available templates with metadata."""
    templates = []
    for fn in os.listdir(TEMPLATES_DIR):
        if fn.endswith(".docx"):
            name = fn[:-5]
            label = name.replace("_", " ").title()
            templates.append({"id": name, "label": label, "file": fn})
    return sorted(templates, key=lambda x: x["label"])
